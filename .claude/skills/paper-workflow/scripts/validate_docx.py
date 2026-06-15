#!/usr/bin/env python3
"""Validate a generated .docx file for basic correctness."""

import re
import sys
from pathlib import Path

HAS_DOCX = False
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    pass


def validate_docx(docx_path: Path) -> dict:
    """Check a docx file: existence, openability, content, headings, tables.

    Returns:
        {
            "valid": bool,
            "errors": [str],
            "warnings": [str],
            "summary": str,
        }
    """
    errors: list[str] = []
    warnings: list[str] = []

    # 1. Existence
    if not docx_path.exists():
        return {
            "valid": False,
            "errors": [f"文件不存在: {docx_path}"],
            "warnings": [],
            "summary": "文件缺失",
        }

    if not HAS_DOCX:
        return {
            "valid": True,
            "errors": [],
            "warnings": ["python-docx 未安装，跳过详细 docx 检查"],
            "summary": "python-docx not available",
        }

    # 2. Openability
    try:
        doc = Document(str(docx_path))
    except Exception as e:
        return {
            "valid": False,
            "errors": [f"无法打开 docx: {e}"],
            "warnings": [],
            "summary": "文件损坏或不可读",
        }

    # 3. Has content
    paragraphs = doc.paragraphs
    if not paragraphs:
        errors.append("文档无段落内容")

    text_content = "\n".join(p.text for p in paragraphs)
    if len(text_content.strip()) < 10:
        warnings.append("文档内容极短，可能为空")

    # 4. Heading styles
    heading_styles = set()
    for p in paragraphs:
        if p.style and p.style.name and p.style.name.startswith("Heading"):
            heading_styles.add(p.style.name)
    if not heading_styles:
        warnings.append("未检测到 Heading 样式")

    # 5. Figure/table numbering
    fig_numbers = re.findall(r"(?:图|Figure|Fig\.?)\s*(\d+)", text_content, re.IGNORECASE)
    if len(set(fig_numbers)) != len(fig_numbers):
        warnings.append("图表编号存在重复")

    tbl_numbers = re.findall(r"(?:表|Table)\s*(\d+)", text_content, re.IGNORECASE)
    if len(set(tbl_numbers)) != len(tbl_numbers):
        warnings.append("表格编号存在重复")

    # 6. Tables
    table_count = len(doc.tables)
    if table_count > 0:
        for i, table in enumerate(doc.tables):
            if not table.rows:
                warnings.append(f"表格 {i + 1} 无行")
    else:
        warnings.append("无表格（非错误，仅提示）")

    return {
        "valid": len(errors) == 0,
        "errors": errors,
        "warnings": warnings,
        "summary": f"段落: {len(paragraphs)}, 标题样式: {len(heading_styles)}, "
                    f"表格: {table_count}",
    }


def main():
    import argparse
    parser = argparse.ArgumentParser(description="Validate a docx file")
    parser.add_argument("docx", help="Path to .docx file")
    args = parser.parse_args()

    result = validate_docx(Path(args.docx))
    print(f"# docx 校验: {args.docx}")
    print(f"  valid: {result['valid']}")
    for e in result["errors"]:
        print(f"  ❌ {e}")
    for w in result["warnings"]:
        print(f"  ⚠ {w}")
    print(f"  {result['summary']}")
    return 0 if result["valid"] else 1


if __name__ == "__main__":
    sys.exit(main())
