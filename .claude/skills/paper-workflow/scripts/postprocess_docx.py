#!/usr/bin/env python3
"""Idempotent Word post-processing for pandoc-generated docx files.

Uses a sidecar JSON file to track processed state.
Sidecar: outputs/qa/<docx-name>.postprocess.json
"""

import hashlib
import json
import sys
from datetime import datetime, timezone
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

POSTPROCESS_VERSION = 1


# ---------------------------------------------------------------------------
# Hashing
# ---------------------------------------------------------------------------

def file_hash(path: Path) -> str:
    """SHA-256 hash of a file."""
    if not path.exists():
        return ""
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(65536), b""):
            h.update(chunk)
    return h.hexdigest()


def profile_hash(profile: dict) -> str:
    """Deterministic hash of a profile dict."""
    raw = json.dumps(profile, sort_keys=True, ensure_ascii=False)
    return hashlib.sha256(raw.encode()).hexdigest()[:16]


# ---------------------------------------------------------------------------
# Sidecar management
# ---------------------------------------------------------------------------

def get_sidecar_path(docx_path: Path, qa_dir: Path) -> Path:
    return qa_dir / f"{docx_path.stem}.postprocess.json"


def load_sidecar(sidecar_path: Path) -> dict | None:
    if not sidecar_path.exists():
        return None
    try:
        with open(sidecar_path, encoding="utf-8") as f:
            return json.load(f)
    except (json.JSONDecodeError, OSError):
        return None


def save_sidecar(sidecar_path: Path, data: dict) -> None:
    sidecar_path.parent.mkdir(parents=True, exist_ok=True)
    with open(sidecar_path, "w", encoding="utf-8") as f:
        json.dump(data, f, indent=2, ensure_ascii=False)


def should_postprocess(docx_path: Path, profile: dict, sidecar_path: Path) -> bool:
    """Check whether postprocessing is needed.

    Returns False if the file hasn't changed since last processing.
    """
    existing = load_sidecar(sidecar_path)
    if existing is None:
        return True

    current_docx_hash = file_hash(docx_path)
    current_profile_hash = profile_hash(profile)

    return (
        existing.get("docx_hash") != current_docx_hash
        or existing.get("profile_hash") != current_profile_hash
        or existing.get("postprocess_version") != POSTPROCESS_VERSION
    )


# ---------------------------------------------------------------------------
# Postprocessing
# ---------------------------------------------------------------------------

def _fix_fonts(doc, warnings: list[str]) -> list[str]:
    """Replace MS Gothic with body font in all runs."""
    changes = []
    for para in doc.paragraphs:
        for run in para.runs:
            font_name = run.font.name
            if font_name and "MS Gothic" in font_name:
                run.font.name = None  # Reset to document default
                changes.append(f"font: MS Gothic → default")
    return changes


def _fix_table_borders(doc, warnings: list[str]) -> list[str]:
    """Ensure tables have basic borders (three-line table style)."""
    changes = []
    for table in doc.tables:
        # Check if table has no border info
        if table.style is None or "Table" in (table.style.name or ""):
            # Set a basic table style
            try:
                table.style = doc.styles["Table"]
                changes.append(f"table: border style applied")
            except KeyError:
                warnings.append("Table style 'Table' not found in document")
    return changes


def postprocess(
    input_path: Path,
    output_path: Path,
    profile: dict,
    qa_dir: Path,
) -> dict:
    """Run postprocessing on a docx file. Idempotent via sidecar.

    Returns:
        {
            "processed": bool,       # True if processing was done
            "skipped": bool,         # True if no processing needed
            "changes": [str],
            "warnings": [str],
            "errors": [str],
        }
    """
    if not HAS_DOCX:
        return {
            "processed": False, "skipped": True,
            "changes": [], "warnings": ["python-docx 未安装，跳过后处理"],
            "errors": [],
        }

    sidecar_path = get_sidecar_path(output_path, qa_dir)
    changes: list[str] = []
    warnings: list[str] = []
    errors: list[str] = []

    # Check if processing is needed
    if not should_postprocess(input_path, profile, sidecar_path):
        return {
            "processed": False, "skipped": True,
            "changes": [],
            "warnings": [],
            "errors": [],
        }

    # Check reference.docx
    ref_doc = profile.get("reference_doc")
    if ref_doc:
        ref_path = Path(ref_doc) if Path(ref_doc).is_absolute() else Path(".claude/skills/paper-workflow") / ref_doc
        if not ref_path.exists():
            warnings.append(f"reference.docx 不存在: {ref_path}。使用默认样式。")

    try:
        doc = Document(str(input_path))
    except Exception as e:
        errors.append(f"无法打开 docx: {e}")
        return {
            "processed": False, "skipped": False,
            "changes": [], "warnings": warnings, "errors": errors,
        }

    # --- Fixes ---

    # 1. Font fixes
    font_changes = _fix_fonts(doc, warnings)
    changes.extend(font_changes)

    # 2. Table borders
    table_changes = _fix_table_borders(doc, warnings)
    changes.extend(table_changes)

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    # Write sidecar
    sidecar_data = {
        "docx_hash": file_hash(output_path),
        "profile_hash": profile_hash(profile),
        "postprocess_version": POSTPROCESS_VERSION,
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "changes": changes,
    }
    save_sidecar(sidecar_path, sidecar_data)

    return {
        "processed": True, "skipped": False,
        "changes": changes, "warnings": warnings, "errors": errors,
    }


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    import argparse
    parser = argparse.ArgumentParser(description="Post-process a docx file")
    parser.add_argument("--input", required=True, help="Input docx path")
    parser.add_argument("--output", required=True, help="Output docx path")
    parser.add_argument("--qa-dir", help="QA directory for sidecar (default: outputs/qa)")
    parser.add_argument("--profile", default="thesis-cn", help="Profile name")
    args = parser.parse_args()

    # Load profile
    import yaml
    PROFILES_DIR = Path(__file__).resolve().parent.parent / "templates" / "profiles"
    profile_path = PROFILES_DIR / f"{args.profile}.yaml"
    profile = {}
    if profile_path.exists():
        with open(profile_path, encoding="utf-8") as f:
            profile = yaml.safe_load(f) or {}

    qa_dir = Path(args.qa_dir) if args.qa_dir else Path(args.output).parent / "qa"

    result = postprocess(Path(args.input), Path(args.output), profile, qa_dir)

    print(json.dumps(result, indent=2, ensure_ascii=False))
    return 1 if result["errors"] else 0


if __name__ == "__main__":
    sys.exit(main())
