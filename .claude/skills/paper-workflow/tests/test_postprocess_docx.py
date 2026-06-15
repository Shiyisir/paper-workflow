"""Test postprocess_docx.py."""

import sys
from pathlib import Path

import pytest

SCRIPTS_DIR = Path(__file__).resolve().parent.parent / "scripts"
sys.path.insert(0, str(SCRIPTS_DIR))

import postprocess_docx as pp


class TestHashing:
    def test_file_hash_consistent(self, tmp_path):
        f = tmp_path / "test.txt"
        f.write_text("hello", encoding="utf-8")
        h1 = pp.file_hash(f)
        h2 = pp.file_hash(f)
        assert h1 == h2

    def test_file_hash_different(self, tmp_path):
        f1 = tmp_path / "a.txt"
        f2 = tmp_path / "b.txt"
        f1.write_text("hello", encoding="utf-8")
        f2.write_text("world", encoding="utf-8")
        assert pp.file_hash(f1) != pp.file_hash(f2)

    def test_file_hash_missing(self):
        assert pp.file_hash(Path("/nonexistent")) == ""

    def test_profile_hash_consistent(self):
        p1 = {"output": "docx", "native_math": "omml"}
        p2 = {"output": "docx", "native_math": "omml"}
        assert pp.profile_hash(p1) == pp.profile_hash(p2)

    def test_profile_hash_different(self):
        p1 = {"output": "docx"}
        p2 = {"output": "tex"}
        assert pp.profile_hash(p1) != pp.profile_hash(p2)


class TestSidecar:
    def test_roundtrip(self, tmp_path):
        sc = tmp_path / "test.postprocess.json"
        data = {"docx_hash": "abc", "profile_hash": "def", "postprocess_version": 1}
        pp.save_sidecar(sc, data)
        loaded = pp.load_sidecar(sc)
        assert loaded == data

    def test_load_nonexistent(self, tmp_path):
        assert pp.load_sidecar(tmp_path / "nonexistent.json") is None

    def test_load_corrupt(self, tmp_path):
        sc = tmp_path / "corrupt.json"
        sc.write_text("not json", encoding="utf-8")
        assert pp.load_sidecar(sc) is None


class TestShouldPostprocess:
    def test_no_sidecar_returns_true(self, tmp_path):
        docx = tmp_path / "test.docx"
        docx.write_text("fake docx", encoding="utf-8")
        assert pp.should_postprocess(docx, {"output": "docx"}, tmp_path / "sidecar.json") is True

    def test_same_hash_returns_false(self, tmp_path):
        docx = tmp_path / "test.docx"
        docx.write_text("fake docx content", encoding="utf-8")
        profile = {"output": "docx", "native_math": "omml"}
        sidecar = tmp_path / "sidecar.json"

        pp.save_sidecar(sidecar, {
            "docx_hash": pp.file_hash(docx),
            "profile_hash": pp.profile_hash(profile),
            "postprocess_version": pp.POSTPROCESS_VERSION,
        })
        assert pp.should_postprocess(docx, profile, sidecar) is False

    def test_changed_docx_returns_true(self, tmp_path):
        docx = tmp_path / "test.docx"
        docx.write_text("old content", encoding="utf-8")
        profile = {"output": "docx"}

        sidecar = tmp_path / "sidecar.json"
        pp.save_sidecar(sidecar, {
            "docx_hash": pp.file_hash(docx),
            "profile_hash": pp.profile_hash(profile),
            "postprocess_version": pp.POSTPROCESS_VERSION,
        })

        # Change the docx
        docx.write_text("new content", encoding="utf-8")
        assert pp.should_postprocess(docx, profile, sidecar) is True


class TestPostprocess:
    def test_idempotent_twice(self, tmp_path):
        """Running postprocess twice on the same input should skip second time."""
        # Create a minimal docx via python-docx
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        docx_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Test content")
        doc.save(str(docx_path))

        qa_dir = tmp_path / "qa"
        profile = {"output": "docx", "native_math": "omml"}
        out_path = tmp_path / "output.docx"

        r1 = pp.postprocess(docx_path, out_path, profile, qa_dir)
        assert r1["processed"] is True or r1["skipped"] is True

        r2 = pp.postprocess(docx_path, out_path, profile, qa_dir)
        assert r2["skipped"] is True  # Should skip second run

    def test_missing_reference_docx_warns(self, tmp_path):
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        docx_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Test")
        doc.save(str(docx_path))

        qa_dir = tmp_path / "qa"
        profile = {"output": "docx", "reference_doc": "nonexistent.docx", "native_math": "omml"}

        result = pp.postprocess(docx_path, tmp_path / "out.docx", profile, qa_dir)
        # Should not crash when reference.docx is missing
        assert "errors" in result

    def test_sidecar_created(self, tmp_path):
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        docx_path = tmp_path / "input.docx"
        doc = Document()
        doc.add_paragraph("Test")
        doc.save(str(docx_path))

        out_path = tmp_path / "output.docx"
        qa_dir = tmp_path / "qa"
        profile = {"output": "docx", "native_math": "omml"}

        result = pp.postprocess(docx_path, out_path, profile, qa_dir)
        if result["processed"]:
            sc_path = pp.get_sidecar_path(out_path, qa_dir)
            assert sc_path.exists()
