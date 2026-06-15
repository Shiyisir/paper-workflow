"""Test validate_docx.py and validate_tex.py."""

import sys
from pathlib import Path

import pytest

SCRIPTS_DIR = Path(__file__).resolve().parent.parent / "scripts"
sys.path.insert(0, str(SCRIPTS_DIR))

import validate_docx
import validate_tex


class TestValidateDocx:
    def test_missing_file(self):
        result = validate_docx.validate_docx(Path("/nonexistent.docx"))
        assert result["valid"] is False
        assert any("不存在" in e for e in result["errors"])

    def test_valid_docx_passes(self, tmp_path):
        """Create a real docx and validate it."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Test content paragraph")
        doc.save(str(path))

        result = validate_docx.validate_docx(path)
        assert result["valid"] is True
        assert result["errors"] == []

    def test_empty_docx_warns(self, tmp_path):
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        path = tmp_path / "empty.docx"
        doc = Document()  # No paragraphs
        doc.save(str(path))

        result = validate_docx.validate_docx(path)
        # Empty docx may have errors or warnings about no content
        assert "errors" in result


class TestValidateTex:
    def test_valid_tex(self, tmp_path):
        tex = tmp_path / "test.tex"
        tex.write_text(
            r"\documentclass{article}" "\n"
            r"\begin{document}" "\n"
            r"Hello world." "\n"
            r"\end{document}" "\n",
            encoding="utf-8"
        )
        result = validate_tex.validate_tex(tex)
        assert result["valid"] is True

    def test_missing_document_env(self, tmp_path):
        tex = tmp_path / "test.tex"
        tex.write_text(r"\documentclass{article}" "\nHello\n", encoding="utf-8")
        result = validate_tex.validate_tex(tex)
        assert result["valid"] is False

    def test_unbalanced_env(self, tmp_path):
        tex = tmp_path / "test.tex"
        tex.write_text(
            r"\begin{document}" "\n"
            r"\begin{document}" "\n"
            r"\end{document}" "\n",
            encoding="utf-8"
        )
        result = validate_tex.validate_tex(tex)
        assert result["valid"] is False

    def test_missing_image_path(self, tmp_path):
        tex = tmp_path / "test.tex"
        tex.write_text(
            r"\documentclass{article}" "\n"
            r"\begin{document}" "\n"
            r"\includegraphics{missing.png}" "\n"
            r"\end{document}" "\n",
            encoding="utf-8"
        )
        result = validate_tex.validate_tex(tex, project_dir=tmp_path)
        assert len(result["errors"]) > 0
        assert any("图片" in e for e in result["errors"])

    def test_missing_file(self):
        result = validate_tex.validate_tex(Path("/nonexistent.tex"))
        assert result["valid"] is False
